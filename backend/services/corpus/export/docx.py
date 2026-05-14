"""DOCX generator — assembles a Word document mirroring the PDF layout.

Visual identity:
  - Same cover composition (brand mark, title, metadata grid, status pill).
  - Same heading hierarchy (Livre / Titre / Chapitre / Section).
  - Same provenance footer (permalink + version + page numbers).

We use python-docx primitives directly rather than a .dotx template — the
content is procedurally generated and the layout is simple enough that
keeping it in code (vs. fighting Word XML) is the lower-friction option.
"""

from __future__ import annotations

import io
from datetime import datetime, timezone

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

from packages.schemas.legal_text import LegalTextRead

from . import _common

# Site identity colors mirrored from the PDF / web app.
NAVY = RGBColor(0x0D, 0x1B, 0x4C)
RED = RGBColor(0xDC, 0x26, 0x26)
RED_LIGHT = RGBColor(0xEF, 0x44, 0x44)
SLATE_700 = RGBColor(0x33, 0x41, 0x55)
SLATE_500 = RGBColor(0x64, 0x74, 0x8B)
SLATE_300 = RGBColor(0xCB, 0xD5, 0xE1)
EMERALD_700 = RGBColor(0x04, 0x78, 0x57)


def render_docx(
    text: LegalTextRead,
    *,
    lang: str = "fr",
    base_url: str = "https://lexhaiti.org",
) -> bytes:
    """Render a legal text as a Word (.docx) document."""
    labels = _common.labels_for(lang)
    tree = _common.build_export_tree(text)
    permalink = f"{base_url.rstrip('/')}/lois/{text.slug}"
    pick = lambda fr, ht: _common.pick_localized(fr, ht, lang) or ""

    doc = Document()

    # Page setup — A4 with comfortable legal-text margins.
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.4)
    section.bottom_margin = Cm(2.4)

    _set_default_font(doc, "Arial", 10.5)
    _build_cover(doc, text, labels, pick, permalink)
    _build_body(doc, text, tree, labels, pick)
    _attach_footer(section, labels, permalink, fmt_version=_common.fmt_date(text.updated_at, lang))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# --------------------------------------------------------------------------- #
# Cover                                                                        #
# --------------------------------------------------------------------------- #

def _build_cover(doc, text, labels, pick, permalink) -> None:
    # Red accent bar at the top of the cover.
    bar = doc.add_paragraph()
    bar.paragraph_format.space_after = Pt(0)
    _set_paragraph_shading(bar, "DC2626")
    run = bar.add_run(" ")
    run.font.size = Pt(2)

    doc.add_paragraph()  # breathing room

    brand = doc.add_paragraph()
    brand.paragraph_format.space_after = Pt(2)
    r1 = brand.add_run("Lex")
    r1.font.size = Pt(20)
    r1.bold = True
    r1.font.color.rgb = NAVY
    r2 = brand.add_run("Haïti")
    r2.font.size = Pt(20)
    r2.bold = True
    r2.font.color.rgb = RED

    tagline = doc.add_paragraph()
    tagline.paragraph_format.space_after = Pt(36)
    tag_run = tagline.add_run(labels.cover_brand_tagline.upper())
    tag_run.font.size = Pt(8)
    tag_run.bold = True
    tag_run.font.color.rgb = SLATE_500

    if text.category:
        kicker = doc.add_paragraph()
        kicker.paragraph_format.space_after = Pt(14)
        kicker_text = labels.category(text.category)
        if text.code_subcategory:
            kicker_text += f" — {text.code_subcategory.replace('_', ' ').capitalize()}"
        kr = kicker.add_run(kicker_text.upper())
        kr.font.size = Pt(9)
        kr.bold = True
        kr.font.color.rgb = RED

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(12)
    tr = title.add_run(pick(text.title_fr, text.title_ht))
    tr.font.name = "Georgia"
    tr.font.size = Pt(26)
    tr.bold = True
    tr.font.color.rgb = NAVY

    subtitle = pick(text.description_fr, text.description_ht)
    if subtitle:
        sub = doc.add_paragraph()
        sub.paragraph_format.space_after = Pt(24)
        sr = sub.add_run(subtitle)
        sr.font.name = "Georgia"
        sr.font.size = Pt(11)
        sr.italic = True
        sr.font.color.rgb = SLATE_700

    # Metadata table — 2 columns × n rows.
    meta_rows = []
    if text.promulgation_date:
        meta_rows.append((labels.cover_promulgation, _common.fmt_date(text.promulgation_date, labels.lang)))
    if text.publication_date:
        meta_rows.append((labels.cover_publication, _common.fmt_date(text.publication_date, labels.lang)))
    if text.moniteur_ref:
        meta_rows.append((labels.cover_moniteur, text.moniteur_ref))
    meta_rows.append((labels.cover_articles, str(len(text.articles))))
    meta_rows.append((labels.cover_status, labels.status(text.status)))

    table = doc.add_table(rows=(len(meta_rows) + 1) // 2, cols=4)
    table.autofit = True
    for idx, (label, value) in enumerate(meta_rows):
        row_idx = idx // 2
        col_offset = (idx % 2) * 2
        label_cell = table.cell(row_idx, col_offset)
        value_cell = table.cell(row_idx, col_offset + 1)
        label_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        value_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

        lp = label_cell.paragraphs[0]
        lp.paragraph_format.space_after = Pt(2)
        lr = lp.add_run(label.upper())
        lr.font.size = Pt(8)
        lr.bold = True
        lr.font.color.rgb = SLATE_500

        vp = value_cell.paragraphs[0]
        vp.paragraph_format.space_after = Pt(2)
        is_status = label == labels.cover_status
        vr = vp.add_run(value)
        vr.font.size = Pt(11)
        vr.bold = True
        if is_status:
            vr.font.color.rgb = _status_color(text.status)
        else:
            vr.font.color.rgb = NAVY

    doc.add_paragraph()  # spacer

    # Cover footer — provenance.
    sep = doc.add_paragraph()
    sep.paragraph_format.space_after = Pt(6)
    sep_run = sep.add_run("─" * 60)
    sep_run.font.size = Pt(6)
    sep_run.font.color.rgb = SLATE_300

    src = doc.add_paragraph()
    src.paragraph_format.space_after = Pt(2)
    sr1 = src.add_run(f"{labels.footer_source} : ")
    sr1.font.size = Pt(8)
    sr1.font.color.rgb = SLATE_500
    sr2 = src.add_run(permalink)
    sr2.font.size = Pt(8)
    sr2.bold = True
    sr2.font.color.rgb = NAVY

    ver = doc.add_paragraph()
    ver_run = ver.add_run(
        f"{labels.footer_version} : {_common.fmt_date(text.updated_at, labels.lang)}  ·  "
        f"{labels.cover_generated} {_common.fmt_date(datetime.now(timezone.utc), labels.lang)}"
    )
    ver_run.font.size = Pt(8)
    ver_run.font.color.rgb = SLATE_500

    # Hard page break to the body.
    pb = doc.add_paragraph().add_run()
    pb.add_break(WD_BREAK.PAGE)


# --------------------------------------------------------------------------- #
# Body                                                                         #
# --------------------------------------------------------------------------- #

def _build_body(doc, text, tree, labels, pick) -> None:
    preamble = pick(text.preamble_fr, text.preamble_ht)
    if preamble:
        h = doc.add_paragraph()
        h.paragraph_format.space_after = Pt(4)
        hr = h.add_run(labels.preamble.upper())
        hr.font.size = Pt(8)
        hr.bold = True
        hr.font.color.rgb = RED

        for para in _common.split_alineas(preamble):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run(para)
            r.italic = True
            r.font.color.rgb = SLATE_700

    for art in tree.orphan_articles:
        _render_article(doc, art, labels, pick)

    for node in tree.roots:
        _render_node(doc, node, labels, pick)


def _render_node(doc, node, labels, pick) -> None:
    h = node.heading
    label = labels.heading(h.level)
    kicker = label + (f" {h.number}" if h.number else "")
    name = pick(h.title_fr, h.title_ht)

    if h.level in {"book", "part"}:
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(2)
    kr = para.add_run(kicker.upper())
    kr.font.size = Pt(9)
    kr.bold = True
    kr.font.color.rgb = RED

    if name:
        title_para = doc.add_paragraph()
        title_para.paragraph_format.space_after = Pt(6)
        nr = title_para.add_run(name)
        nr.font.name = "Georgia"
        nr.bold = True
        nr.font.color.rgb = NAVY
        nr.font.size = _heading_size(h.level)

    for art in node.articles:
        _render_article(doc, art, labels, pick)

    for child in node.children:
        _render_node(doc, child, labels, pick)


def _heading_size(level: str) -> Pt:
    return {
        "book": Pt(16),
        "part": Pt(16),
        "title": Pt(14),
        "chapter": Pt(12.5),
        "section": Pt(11),
        "subsection": Pt(11),
    }.get(level, Pt(11))


def _render_article(doc, art, labels, pick) -> None:
    label = doc.add_paragraph()
    label.paragraph_format.space_before = Pt(6)
    label.paragraph_format.space_after = Pt(2)
    lr = label.add_run(_common.article_label(art, labels))
    lr.font.size = Pt(9.5)
    lr.bold = True
    lr.font.color.rgb = RED

    title = pick(art.title_fr, art.title_ht)
    if title:
        tp = doc.add_paragraph()
        tp.paragraph_format.space_after = Pt(2)
        tr = tp.add_run(title)
        tr.font.name = "Georgia"
        tr.bold = True
        tr.font.color.rgb = NAVY
        tr.font.size = Pt(11)

    body = pick(art.content_fr, art.content_ht)
    if body:
        for para in _common.split_alineas(body):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(2)
            if _common.is_list_item(para):
                # Hanging indent: marker (a)/b)/1°) sits at left, continuation
                # lines align to the marker text — same effect as the PDF.
                p.paragraph_format.left_indent = Cm(0.6)
                p.paragraph_format.first_line_indent = Cm(-0.6)
            p.add_run(para)

    if labels.lang == "ht" and not art.content_ht and art.content_fr:
        notice = doc.add_paragraph()
        nr = notice.add_run(labels.no_haitian_translation)
        nr.font.size = Pt(8)
        nr.italic = True
        nr.font.color.rgb = SLATE_500


# --------------------------------------------------------------------------- #
# Footer                                                                       #
# --------------------------------------------------------------------------- #

def _attach_footer(section, labels, permalink: str, fmt_version: str) -> None:
    """Add a provenance footer with permalink + version on the left, page X / Y on the right.

    The cover page footer is suppressed by setting different_first_page so the
    cover keeps its own formatting and isn't doubled-up.
    """
    section.different_first_page_header_footer = True

    footer = section.footer
    footer.is_linked_to_previous = False

    para = footer.paragraphs[0]
    para.paragraph_format.tab_stops.add_tab_stop(section.page_width - section.left_margin - section.right_margin)

    left = para.add_run(
        f"{labels.footer_source} : {permalink}  ·  "
        f"{labels.footer_version} {fmt_version}"
    )
    left.font.size = Pt(7.5)
    left.font.color.rgb = SLATE_500

    para.add_run("\t")

    # Page number field — Word renders this dynamically.
    _add_page_field(para, "PAGE")
    para.add_run(" / ").font.size = Pt(7.5)
    _add_page_field(para, "NUMPAGES")


def _add_page_field(paragraph, instr: str) -> None:
    run = paragraph.add_run()
    run.font.size = Pt(7.5)
    run.font.color.rgb = SLATE_500

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._element.append(fld_begin)

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instr
    run._element.append(instr_text)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._element.append(fld_end)


# --------------------------------------------------------------------------- #
# Helpers                                                                      #
# --------------------------------------------------------------------------- #

def _set_default_font(doc, name: str, size_pt: float) -> None:
    style = doc.styles["Normal"]
    style.font.name = name
    style.font.size = Pt(size_pt)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:cs"), name)


def _status_color(status) -> RGBColor:
    """Color the cover-page status text by its meaning — green for in-force,
    red for abrogated, amber for suspended."""
    s = str(status)
    if s == "abrogated":
        return RGBColor(0xB9, 0x1C, 0x1C)
    if s == "suspended":
        return RGBColor(0xB4, 0x53, 0x09)
    return EMERALD_700


def _set_paragraph_shading(paragraph, hex_rgb: str) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_rgb)
    pPr.append(shd)
